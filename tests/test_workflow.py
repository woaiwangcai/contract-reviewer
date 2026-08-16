from pathlib import Path

import pytest
from docx import Document

import src.workflow as workflow_module
from src.workflow import ReviewWorkflowError, _safe_stem, run_review


def test_workflow_writes_markdown_and_html_without_api(tmp_path: Path) -> None:
    input_path = tmp_path / "示例:合同.docx"
    document = Document()
    document.add_heading("服务合同", level=1)
    document.add_paragraph("合同金额为 100 元。")
    document.save(input_path)

    skill_path = tmp_path / "skill.md"
    skill_path.write_text("审查输入合同。", encoding="utf-8")
    template_path = tmp_path / "template.html"
    template_path.write_text("<title>$title</title><main>$content</main>", encoding="utf-8")

    captured: dict[str, str] = {}

    def fake_reviewer(skill: str, contract: str) -> str:
        captured["skill"] = skill
        captured["contract"] = contract
        return "# 合同审查结果\n\n无重大风险。\n"

    artifacts = run_review(
        input_path=input_path,
        skill_path=skill_path,
        output_dir=tmp_path / "output",
        template_path=template_path,
        reviewer=fake_reviewer,
    )

    assert captured["skill"] == "审查输入合同。"
    assert "合同金额为 100 元" in captured["contract"]
    assert artifacts.markdown_path.name == "示例_合同-review.md"
    assert artifacts.markdown_path.is_file()
    assert artifacts.html_path.is_file()
    assert "合同审查结果" in artifacts.html_path.read_text(encoding="utf-8")


def test_safe_stem_falls_back_for_invalid_name() -> None:
    assert _safe_stem("<>:\\|?*") == "contract"


def test_existing_reports_receive_a_sequence_suffix(tmp_path: Path) -> None:
    input_path = tmp_path / "contract.docx"
    document = Document()
    document.add_paragraph("合同正文")
    document.save(input_path)
    skill_path = tmp_path / "skill.md"
    skill_path.write_text("审查输入合同。", encoding="utf-8")
    template_path = tmp_path / "template.html"
    template_path.write_text("<main>$content</main>", encoding="utf-8")
    output_dir = tmp_path / "output"
    reviewer = lambda _skill, _contract: "# 审查结果\n"

    first = run_review(input_path, skill_path, output_dir, template_path, reviewer)
    second = run_review(input_path, skill_path, output_dir, template_path, reviewer)

    assert first.markdown_path.name == "contract-review.md"
    assert second.markdown_path.name == "contract-review-2.md"


def test_failed_second_commit_removes_partial_report(tmp_path: Path, monkeypatch) -> None:
    input_path = tmp_path / "contract.docx"
    document = Document()
    document.add_paragraph("合同正文")
    document.save(input_path)
    skill_path = tmp_path / "skill.md"
    skill_path.write_text("审查输入合同。", encoding="utf-8")
    template_path = tmp_path / "template.html"
    template_path.write_text("<main>$content</main>", encoding="utf-8")
    output_dir = tmp_path / "output"

    original_replace = workflow_module.os.replace
    calls = 0

    def fail_second_replace(source, destination) -> None:
        nonlocal calls
        calls += 1
        if calls == 2:
            raise OSError("simulated write failure")
        original_replace(source, destination)

    monkeypatch.setattr(workflow_module.os, "replace", fail_second_replace)

    with pytest.raises(ReviewWorkflowError, match="无法写入"):
        run_review(
            input_path,
            skill_path,
            output_dir,
            template_path,
            lambda _skill, _contract: "# 审查结果\n",
        )

    assert not list(output_dir.glob("contract-review*"))
