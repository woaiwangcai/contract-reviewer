from pathlib import Path
from zipfile import ZIP_DEFLATED, ZipFile

import pytest
from docx import Document
from docx.oxml.ns import qn
from lxml import etree

import src.docx_to_md as docx_module
from src.docx_to_md import DocxConversionError, docx_to_markdown


def _rewrite_document_xml(path: Path, transform) -> None:
    with ZipFile(path) as source:
        entries = {item.filename: source.read(item.filename) for item in source.infolist()}

    root = etree.fromstring(entries["word/document.xml"])
    transform(root)
    entries["word/document.xml"] = etree.tostring(
        root,
        xml_declaration=True,
        encoding="UTF-8",
        standalone=True,
    )

    rebuilt = path.with_suffix(".rebuilt.docx")
    with ZipFile(rebuilt, "w", ZIP_DEFLATED) as target:
        for name, content in entries.items():
            target.writestr(name, content)
    rebuilt.replace(path)


def test_converts_headings_paragraphs_lists_and_tables(tmp_path: Path) -> None:
    path = tmp_path / "sample.docx"
    document = Document()
    document.add_heading("示例合同", level=1)
    document.add_paragraph("这是合同正文。")
    document.add_paragraph("第一项", style="List Bullet")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "项目"
    table.cell(0, 1).text = "内容"
    table.cell(1, 0).text = "金额"
    table.cell(1, 1).text = "100 元"
    document.save(path)

    result = docx_to_markdown(path)

    assert "# 示例合同" in result
    assert "这是合同正文。" in result
    assert "- 第一项" in result
    assert "| 项目 | 内容 |" in result
    assert "| 金额 | 100 元 |" in result


def test_rejects_non_docx_file(tmp_path: Path) -> None:
    path = tmp_path / "contract.txt"
    path.write_text("not a docx", encoding="utf-8")

    with pytest.raises(DocxConversionError, match="仅支持"):
        docx_to_markdown(path)


def test_includes_content_controls_and_tracked_insertions_but_not_deletions(
    tmp_path: Path,
) -> None:
    path = tmp_path / "tracked.docx"
    document = Document()
    document.add_paragraph("CONTROL_SENTINEL")
    document.add_paragraph("CURRENT_SENTINEL")
    document.save(path)

    def transform(root) -> None:
        namespaces = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
        control_text = root.xpath(".//w:t[text()='CONTROL_SENTINEL']", namespaces=namespaces)[0]
        control_paragraph = control_text.getparent().getparent()
        parent = control_paragraph.getparent()
        index = parent.index(control_paragraph)
        parent.remove(control_paragraph)
        content_control = etree.Element(qn("w:sdt"))
        content_control.append(etree.Element(qn("w:sdtPr")))
        content = etree.SubElement(content_control, qn("w:sdtContent"))
        content.append(control_paragraph)
        parent.insert(index, content_control)

        current_text = root.xpath(".//w:t[text()='CURRENT_SENTINEL']", namespaces=namespaces)[0]
        current_run = current_text.getparent()
        paragraph = current_run.getparent()
        paragraph.remove(current_run)
        insertion = etree.SubElement(paragraph, qn("w:ins"))
        insertion.set(qn("w:id"), "1")
        insertion.append(current_run)
        deletion = etree.SubElement(paragraph, qn("w:del"))
        deletion.set(qn("w:id"), "2")
        deleted_run = etree.SubElement(deletion, qn("w:r"))
        deleted_text = etree.SubElement(deleted_run, qn("w:delText"))
        deleted_text.text = "DELETED_SENTINEL"

    _rewrite_document_xml(path, transform)

    result = docx_to_markdown(path)

    assert "CONTROL_SENTINEL" in result
    assert "CURRENT_SENTINEL" in result
    assert "DELETED_SENTINEL" not in result


def test_malformed_internal_xml_is_reported_as_conversion_error(tmp_path: Path) -> None:
    path = tmp_path / "broken.docx"
    document = Document()
    document.add_paragraph("text")
    document.save(path)

    with ZipFile(path) as source:
        entries = {item.filename: source.read(item.filename) for item in source.infolist()}
    entries["word/document.xml"] = b"<w:document"
    with ZipFile(path, "w", ZIP_DEFLATED) as target:
        for name, content in entries.items():
            target.writestr(name, content)

    with pytest.raises(DocxConversionError, match="损坏"):
        docx_to_markdown(path)


def test_rejects_archive_over_uncompressed_limit(tmp_path: Path, monkeypatch) -> None:
    path = tmp_path / "large.docx"
    document = Document()
    document.add_paragraph("text")
    document.save(path)
    monkeypatch.setattr(docx_module, "MAX_TOTAL_UNCOMPRESSED_BYTES", 1)

    with pytest.raises(DocxConversionError, match="解压后"):
        docx_to_markdown(path)


def test_merged_cells_are_not_duplicated(tmp_path: Path) -> None:
    path = tmp_path / "merged.docx"
    document = Document()
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).merge(table.cell(0, 1)).text = "合并标题"
    table.cell(1, 0).text = "A"
    table.cell(1, 1).text = "B"
    document.save(path)

    result = docx_to_markdown(path)

    assert "| 合并标题 |  |" in result
    assert "| 合并标题 | 合并标题 |" not in result


def test_includes_header_and_footer_text(tmp_path: Path) -> None:
    path = tmp_path / "stories.docx"
    document = Document()
    document.add_paragraph("合同正文")
    document.sections[0].header.paragraphs[0].text = "页眉中的合同编号"
    document.sections[0].footer.paragraphs[0].text = "页脚中的补充说明"
    document.save(path)

    result = docx_to_markdown(path)

    assert "## 页眉" in result
    assert "页眉中的合同编号" in result
    assert "## 页脚" in result
    assert "页脚中的补充说明" in result
