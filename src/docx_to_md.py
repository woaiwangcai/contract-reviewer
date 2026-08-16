from __future__ import annotations

import re
from pathlib import Path
from zipfile import BadZipFile, ZipFile

from docx import Document
from docx.document import Document as DocumentObject
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph


class DocxConversionError(RuntimeError):
    """Raised when a DOCX file cannot be converted safely."""


MAX_INPUT_BYTES = 25 * 1024 * 1024
MAX_ARCHIVE_ENTRIES = 2_000
MAX_TOTAL_UNCOMPRESSED_BYTES = 100 * 1024 * 1024
MAX_SINGLE_PART_BYTES = 30 * 1024 * 1024


def _iter_blocks(document: DocumentObject):
    def walk(parent):
        for child in parent.iterchildren():
            if isinstance(child, CT_P):
                yield Paragraph(child, document)
            elif isinstance(child, CT_Tbl):
                yield Table(child, document)
            else:
                yield from walk(child)

    yield from walk(document.element.body)


def _current_text(element) -> str:
    parts: list[str] = []

    def walk(node) -> None:
        if node.tag == qn("w:del"):
            return
        if node.tag == qn("w:t") and node.text:
            parts.append(node.text)
            return
        if node.tag == qn("w:tab"):
            parts.append("\t")
            return
        if node.tag in {qn("w:br"), qn("w:cr")}:
            parts.append("\n")
            return
        for child in node.iterchildren():
            walk(child)

    walk(element)
    return "".join(parts)


def _heading_level(style_name: str) -> int | None:
    match = re.search(r"(?:Heading|标题)\s*([1-6])", style_name, re.IGNORECASE)
    return int(match.group(1)) if match else None


def _paragraph_to_markdown(paragraph: Paragraph) -> str:
    text = _current_text(paragraph._p).strip()
    if not text:
        return ""

    style_name = paragraph.style.name if paragraph.style else ""
    heading_level = _heading_level(style_name)
    if heading_level:
        return f"{'#' * heading_level} {text}"

    has_numbering = paragraph._p.pPr is not None and paragraph._p.pPr.numPr is not None
    lowered_style = style_name.lower()
    if "bullet" in lowered_style or "项目符号" in style_name:
        return f"- {text}"
    if has_numbering or "number" in lowered_style or "编号" in style_name:
        return f"1. {text}"
    return text


def _escape_cell(value: str) -> str:
    compact = "<br>".join(part.strip() for part in value.splitlines() if part.strip())
    return compact.replace("\\", "\\\\").replace("|", "\\|")


def _table_to_markdown(table: Table) -> str:
    rows: list[list[str]] = []
    for row in table.rows:
        seen_cells: set[int] = set()
        values: list[str] = []
        for cell in row.cells:
            cell_id = id(cell._tc)
            value = "" if cell_id in seen_cells else _current_text(cell._tc)
            seen_cells.add(cell_id)
            values.append(_escape_cell(value))
        rows.append(values)
    if not rows:
        return ""

    width = max(len(row) for row in rows)
    normalized = [row + [""] * (width - len(row)) for row in rows]
    header = normalized[0]
    lines = [
        f"| {' | '.join(header)} |",
        f"| {' | '.join(['---'] * width)} |",
    ]
    lines.extend(f"| {' | '.join(row)} |" for row in normalized[1:])
    return "\n".join(lines)


def _validate_docx_package(path: Path) -> None:
    try:
        if path.stat().st_size > MAX_INPUT_BYTES:
            raise DocxConversionError("DOCX 文件超过 25 MB 限制。")
        with ZipFile(path) as archive:
            entries = archive.infolist()
            if len(entries) > MAX_ARCHIVE_ENTRIES:
                raise DocxConversionError("DOCX 包含过多内部文件。")
            if any(entry.flag_bits & 0x1 for entry in entries):
                raise DocxConversionError("不支持加密的 DOCX 文件。")
            if any(entry.file_size > MAX_SINGLE_PART_BYTES for entry in entries):
                raise DocxConversionError("DOCX 内部文件超过安全大小限制。")
            total_size = sum(entry.file_size for entry in entries)
            if total_size > MAX_TOTAL_UNCOMPRESSED_BYTES:
                raise DocxConversionError("DOCX 解压后超过 100 MB 安全限制。")
            for entry in entries:
                if (
                    entry.file_size > 5 * 1024 * 1024
                    and entry.compress_size > 0
                    and entry.file_size / entry.compress_size > 200
                ):
                    raise DocxConversionError("DOCX 压缩比异常，已停止解析。")
    except BadZipFile as exc:
        raise DocxConversionError("DOCX 文件损坏或格式无效。") from exc
    except OSError as exc:
        raise DocxConversionError(f"无法读取 DOCX 文件：{exc}") from exc


def _story_markdown(document: DocumentObject) -> list[str]:
    collected: list[str] = []
    seen: set[tuple[str, str]] = set()

    for section in document.sections:
        stories = [("页眉", section.header), ("页脚", section.footer)]
        if section.different_first_page_header_footer:
            stories.extend(
                [("首页页眉", section.first_page_header), ("首页页脚", section.first_page_footer)]
            )
        if document.settings.odd_and_even_pages_header_footer:
            stories.extend(
                [("偶数页页眉", section.even_page_header), ("偶数页页脚", section.even_page_footer)]
            )

        for label, story in stories:
            parts = [
                _current_text(paragraph._p).strip()
                for paragraph in story.paragraphs
                if _current_text(paragraph._p).strip()
            ]
            parts.extend(
                rendered
                for table in story.tables
                if (rendered := _table_to_markdown(table))
            )
            content = "\n\n".join(parts).strip()
            normalized = re.sub(r"\s+", "", content)
            if re.fullmatch(r"第(?:\d+)?页(?:/共\d+页)?", normalized):
                continue
            key = (label, content)
            if content and key not in seen:
                seen.add(key)
                collected.append(f"## {label}\n\n{content}")

    return collected


def docx_to_markdown(path: Path) -> str:
    path = Path(path)
    if path.suffix.lower() != ".docx":
        raise DocxConversionError("仅支持 .docx 文件。")
    if not path.is_file():
        raise DocxConversionError(f"找不到输入文件：{path}")

    _validate_docx_package(path)

    try:
        document = Document(path)
        blocks: list[str] = []
        for block in _iter_blocks(document):
            rendered = (
                _paragraph_to_markdown(block)
                if isinstance(block, Paragraph)
                else _table_to_markdown(block)
            )
            if rendered:
                blocks.append(rendered)
        blocks.extend(_story_markdown(document))
    except OSError as exc:
        raise DocxConversionError(f"无法读取 DOCX 文件：{exc}") from exc
    except Exception as exc:
        raise DocxConversionError("DOCX 文件损坏、结构异常或包含不支持的内容。") from exc

    markdown = "\n\n".join(blocks).strip()
    if not markdown:
        raise DocxConversionError("DOCX 中没有可审查的文本。")
    return markdown + "\n"
