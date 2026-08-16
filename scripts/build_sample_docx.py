from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


INK = "17211B"
BLUE = "2E74B5"
MUTED = "5E6862"
TABLE_FILL = "F2F4F7"
CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_MARGIN_DXA = {"top": 80, "bottom": 80, "start": 120, "end": 120}


def set_run_font(run, *, size: float = 11, bold: bool = False, color: str = INK) -> None:
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")


def configure_styles(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    tokens = {
        "Heading 1": (16, BLUE, 16, 8),
        "Heading 2": (13, BLUE, 12, 6),
        "Heading 3": (12, "1F4D78", 8, 4),
    }
    for style_name, (size, color, before, after) in tokens.items():
        style = document.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.10
        style.paragraph_format.keep_with_next = True


def set_cell_margins(cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in CELL_MARGIN_DXA.items():
        element = tc_mar.find(qn(f"w:{side}"))
        if element is None:
            element = OxmlElement(f"w:{side}")
            tc_mar.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[int]) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr

    tbl_width = tbl_pr.first_child_found_in("w:tblW")
    if tbl_width is None:
        tbl_width = OxmlElement("w:tblW")
        tbl_pr.append(tbl_width)
    tbl_width.set(qn("w:w"), str(sum(widths)))
    tbl_width.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for cell, width in zip(row.cells, widths, strict=True):
            tc_w = cell._tc.get_or_add_tcPr().get_or_add_tcW()
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("第 ")
    set_run_font(run, size=9, color=MUTED)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, end])
    tail = paragraph.add_run(" 页")
    set_run_font(tail, size=9, color=MUTED)


def add_clause(document: Document, title: str, paragraphs: list[str]) -> None:
    document.add_heading(title, level=2)
    for text in paragraphs:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.first_line_indent = Inches(0.29)
        paragraph.add_run(text)


def build(output_path: Path) -> None:
    document = Document()
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    configure_styles(document)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header_run = header.add_run("虚构示例｜仅用于项目演示")
    set_run_font(header_run, size=9, color=MUTED)
    add_page_number(section.footer.paragraphs[0])

    title = document.add_heading("数据分析平台技术服务合同（虚构示例）", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(18)
    for run in title.runs:
        set_run_font(run, size=20, bold=True, color=INK)

    notice = document.add_paragraph()
    notice.alignment = WD_ALIGN_PARAGRAPH.CENTER
    notice.paragraph_format.space_after = Pt(18)
    notice_run = notice.add_run("本文件中的主体、金额、条款和签署信息均为虚构内容")
    set_run_font(notice_run, size=10, color=MUTED)

    table = document.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    rows = [
        ("项目", "内容"),
        ("甲方", "澄明科技（示例）有限公司"),
        ("乙方", "远岭数据（示例）有限公司"),
        ("合同金额", "人民币 120,000 元（含税）"),
        ("服务期限", "平台交付后一年"),
    ]
    for index, values in enumerate(rows):
        cells = table.rows[0].cells if index == 0 else table.add_row().cells
        for cell, value in zip(cells, values, strict=True):
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            run = paragraph.add_run(value)
            set_run_font(run, size=10.5, bold=index == 0)
            if index == 0:
                shade_cell(cell, TABLE_FILL)
    set_table_geometry(table, [1700, CONTENT_WIDTH_DXA - 1700])

    intro = document.add_paragraph()
    intro.paragraph_format.space_before = Pt(14)
    intro.paragraph_format.first_line_indent = Inches(0.29)
    intro.add_run("甲乙双方经协商，就数据分析平台开发、交付及技术服务事宜订立本合同。")

    add_clause(document, "第一条 服务内容", [
        "乙方按照附件《功能清单》为甲方开发并部署数据分析平台，提供数据看板、报表导出和账号权限管理功能。附件由双方另行确认。",
        "乙方负责平台安装、基础培训和上线支持。超出附件范围的需求由双方另行协商。",
    ])
    add_clause(document, "第二条 费用与付款", [
        "合同总价为人民币 120,000 元（含税）。合同生效后五个工作日内，甲方向乙方支付合同总价的 70%；平台验收后五个工作日内支付剩余 30%。",
        "乙方应在收款前向甲方开具合法有效的增值税发票。",
    ])
    add_clause(document, "第三条 交付与验收", [
        "乙方应在收到首付款后 30 日内完成平台交付。甲方应在收到交付通知后三日内提出书面异议，逾期未提出异议视为验收合格。",
        "甲方提出异议的，乙方应在合理期限内完成修改，具体整改期限和复验方式由双方另行协商。",
    ])
    add_clause(document, "第四条 技术服务", [
        "自平台验收之日起一年内，乙方提供工作日 9:00 至 18:00 的远程技术支持。服务期满后的维护费用由双方另行协商。",
    ])
    add_clause(document, "第五条 数据使用", [
        "甲方授权乙方处理履行本合同所需的数据。乙方有权将甲方数据用于产品优化及其他商业目的，并可根据业务需要委托第三方处理。",
    ])
    add_clause(document, "第六条 保密", [
        "双方对履约过程中获悉的商业秘密承担保密义务，保密期限为合同终止后两年。法律法规或有权机关要求披露的除外。",
    ])
    add_clause(document, "第七条 责任限制", [
        "一方违约应赔偿对方因此遭受的直接损失。无论基于何种原因，乙方在本合同项下的累计赔偿责任不超过乙方已收服务费的 10%。",
    ])
    add_clause(document, "第八条 解除", [
        "任何一方严重违约且在收到书面通知后十日内未改正的，守约方有权解除合同。合同解除后，甲方已付款项不予退还。",
    ])
    add_clause(document, "第九条 争议解决", [
        "因本合同产生的争议，双方应先协商解决；协商不成的，任何一方可向被告所在地有管辖权的人民法院提起诉讼。",
    ])
    add_clause(document, "第十条 生效与文本", [
        "本合同自双方盖章之日起生效，一式两份，双方各执一份，具有同等效力。",
    ])

    document.add_paragraph()
    signature = document.add_table(rows=2, cols=2)
    signature.style = "Table Grid"
    signature_data = [
        ("甲方（盖章）：", "乙方（盖章）："),
        ("日期：____年__月__日", "日期：____年__月__日"),
    ]
    for row, values in zip(signature.rows, signature_data, strict=True):
        for cell, value in zip(row.cells, values, strict=True):
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(12)
            run = paragraph.add_run(value)
            set_run_font(run)
    set_table_geometry(signature, [CONTENT_WIDTH_DXA // 2, CONTENT_WIDTH_DXA // 2])

    document.core_properties.title = "虚构数据分析平台技术服务合同"
    document.core_properties.subject = "Contract Reviewer fictional example"
    document.core_properties.creator = "Contract Reviewer"
    document.core_properties.last_modified_by = "Contract Reviewer"
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


if __name__ == "__main__":
    destination = Path(sys.argv[1]) if len(sys.argv) > 1 else Path("examples/sample_contract.raw.docx")
    build(destination)
